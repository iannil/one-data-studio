"""
文档解析器
支持PDF、Word、Excel、图片等多种格式的解析
"""

import os
import io
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from enum import Enum

import numpy as np

logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """文档格式枚举"""
    PDF = "pdf"
    PDF_SCANNED = "scanned_pdf"
    IMAGE = "image"
    WORD = "word"
    EXCEL = "excel"
    UNKNOWN = "unknown"


class DocumentParser:
    """文档解析器"""

    def __init__(self, temp_dir: str = "/tmp/ocr"):
        self.temp_dir = temp_dir
        os.makedirs(temp_dir, exist_ok=True)

        # 检测可用的库
        self._init_dependencies()

    def _init_dependencies(self):
        """初始化依赖库"""
        self._has_pdf2image = False
        self._has_pymupdf = False
        self._has_pdfplumber = False
        self._has_pillow = False
        self._has_python_docx = False
        self._has_openpyxl = False
        self._has_camelot = False

        # PDF处理
        try:
            import pdf2image
            self._has_pdf2image = True
        except ImportError:
            pass

        try:
            import fitz  # PyMuPDF
            self._pymupdf = fitz
            self._has_pymupdf = True
        except ImportError:
            pass

        try:
            import pdfplumber
            self._pdfplumber = pdfplumber
            self._has_pdfplumber = True
        except ImportError:
            pass

        # 图片处理
        try:
            from PIL import Image
            self._PIL = Image
            self._has_pillow = True
        except ImportError:
            pass

        # Word处理
        try:
            import docx
            self._python_docx = docx
            self._has_python_docx = True
        except ImportError:
            pass

        # Excel处理
        try:
            import openpyxl
            self._openpyxl = openpyxl
            self._has_openpyxl = True
        except ImportError:
            pass

        # 表格提取
        try:
            import camelot
            self._camelot = camelot
            self._has_camelot = True
        except ImportError:
            pass

        logger.info(f"DocumentParser dependencies: "
                   f"pdf2image={self._has_pdf2image}, "
                   f"PyMuPDF={self._has_pymupdf}, "
                   f"pdfplumber={self._has_pdfplumber}, "
                   f"PIL={self._has_pillow}, "
                   f"docx={self._has_python_docx}, "
                   f"openpyxl={self._has_openpyxl}, "
                   f"camelot={self._has_camelot}")

    def detect_format(self, file_path: str, file_content: bytes = None) -> DocumentFormat:
        """检测文档格式"""
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            # 检测是否为扫描件
            if file_content:
                return self._detect_scanned_pdf(file_content)
            return DocumentFormat.PDF
        elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp']:
            return DocumentFormat.IMAGE
        elif ext in ['.docx', '.doc']:
            return DocumentFormat.WORD
        elif ext in ['.xlsx', '.xls']:
            return DocumentFormat.EXCEL
        else:
            return DocumentFormat.UNKNOWN

    def _detect_scanned_pdf(self, file_content: bytes) -> DocumentFormat:
        """检测PDF是否为扫描件"""
        if not self._has_pymupdf:
            return DocumentFormat.PDF

        try:
            doc = self._pymupdf.open(stream=file_content)
            page = doc[0]

            # 获取页面文本
            text = page.get_text()

            # 如果文本很少，可能是扫描件
            if len(text.strip()) < 50:
                return DocumentFormat.PDF_SCANNED

            return DocumentFormat.PDF
        except Exception as e:
            logger.error(f"Error detecting scanned PDF: {e}")
            return DocumentFormat.PDF

    def parse(self, file_path: str, file_content: bytes = None) -> Dict:
        """
        解析文档
        返回: {
            "format": "pdf",
            "pages": [
                {
                    "number": 1,
                    "image": np.ndarray,
                    "text": "page text",
                    "tables": [...]
                }
            ],
            "metadata": {...}
        }
        """
        file_format = self.detect_format(file_path, file_content)

        result = {
            "format": file_format.value,
            "pages": [],
            "metadata": {}
        }

        if file_format == DocumentFormat.PDF or file_format == DocumentFormat.PDF_SCANNED:
            result = self._parse_pdf(file_path, file_content, file_format)
        elif file_format == DocumentFormat.IMAGE:
            result = self._parse_image(file_path, file_content)
        elif file_format == DocumentFormat.WORD:
            result = self._parse_word(file_path, file_content)
        elif file_format == DocumentFormat.EXCEL:
            result = self._parse_excel(file_path, file_content)
        else:
            raise ValueError(f"Unsupported document format: {file_format}")

        return result

    def _parse_pdf(self, file_path: str, file_content: bytes, file_format: DocumentFormat) -> Dict:
        """解析PDF文档"""
        result = {
            "format": file_format.value,
            "pages": [],
            "metadata": {}
        }

        # 保存临时文件
        temp_path = self._save_temp_file(file_path, file_content)

        try:
            if file_format == DocumentFormat.PDF_SCANNED:
                # 扫描件需要转换为图片
                pages = self._pdf_to_images(temp_path)
                for i, page_image in enumerate(pages):
                    result["pages"].append({
                        "number": i + 1,
                        "image": page_image,
                        "text": "",
                        "tables": []
                    })
            else:
                # 普通PDF可以提取文本
                if self._has_pymupdf:
                    result = self._parse_pdf_with_pymupdf(temp_path, file_format)
                elif self._has_pdfplumber:
                    result = self._parse_pdf_with_pdfplumber(temp_path, file_format)
                else:
                    # 回退到图片转换
                    pages = self._pdf_to_images(temp_path)
                    for i, page_image in enumerate(pages):
                        result["pages"].append({
                            "number": i + 1,
                            "image": page_image,
                            "text": "",
                            "tables": []
                        })

        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)

        return result

    def _parse_pdf_with_pymupdf(self, file_path: str, file_format: DocumentFormat) -> Dict:
        """使用PyMuPDF解析PDF"""
        result = {
            "format": file_format.value,
            "pages": [],
            "metadata": {}
        }

        doc = self._pymupdf.open(file_path)

        # 元数据
        result["metadata"] = {
            "page_count": doc.page_count,
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", "")
        }

        for page_num in range(doc.page_count):
            page = doc[page_num]

            # 提取文本
            text = page.get_text()

            # 转换为图片用于OCR
            pix = page.get_pixmap(matrix=self._pymupdf.Matrix(2, 2))  # 2x缩放提高清晰度
            img_data = pix.tobytes("png")
            from PIL import Image
            import io
            page_image = np.array(Image.open(io.BytesIO(img_data)))

            # 提取表格
            tables = self._extract_tables_from_page(page)

            result["pages"].append({
                "number": page_num + 1,
                "image": page_image,
                "text": text,
                "tables": tables
            })

        doc.close()
        return result

    def _parse_pdf_with_pdfplumber(self, file_path: str, file_format: DocumentFormat) -> Dict:
        """使用pdfplumber解析PDF"""
        result = {
            "format": file_format.value,
            "pages": [],
            "metadata": {}
        }

        with self._pdfplumber.open(file_path) as pdf:
            result["metadata"]["page_count"] = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages):
                # 提取文本
                text = page.extract_text() or ""

                # 提取表格
                tables = page.extract_tables() or []

                # 转换页面为图片
                page_image = None
                if self._has_pdf2image:
                    from pdf2image import convert_from_path
                    images = convert_from_path(file_path, first_page=page_num + 1, last_page=page_num + 1)
                    if images:
                        page_image = np.array(images[0])

                result["pages"].append({
                    "number": page_num + 1,
                    "image": page_image,
                    "text": text,
                    "tables": tables
                })

        return result

    def _pdf_to_images(self, file_path: str) -> List[np.ndarray]:
        """将PDF转换为图片列表"""
        images = []

        if self._has_pdf2image:
            from pdf2image import convert_from_path
            pil_images = convert_from_path(file_path, dpi=200)
            images = [np.array(img) for img in pil_images]
        elif self._has_pymupdf:
            doc = self._pymupdf.open(file_path)
            for page in doc:
                pix = page.get_pixmap(matrix=self._pymupdf.Matrix(2, 2))
                img_data = pix.tobytes("png")
                from PIL import Image
                import io
                page_image = np.array(Image.open(io.BytesIO(img_data)))
                images.append(page_image)
            doc.close()

        return images

    def _parse_image(self, file_path: str, file_content: bytes) -> Dict:
        """解析图片"""
        result = {
            "format": DocumentFormat.IMAGE.value,
            "pages": [],
            "metadata": {}
        }

        if file_content:
            from PIL import Image
            import io
            image = np.array(Image.open(io.BytesIO(file_content)))
        else:
            image = np.array(self._PIL.open(file_path))

        result["pages"].append({
            "number": 1,
            "image": image,
            "text": "",
            "tables": []
        })

        result["metadata"] = {
            "width": image.shape[1],
            "height": image.shape[0],
            "channels": image.shape[2] if len(image.shape) > 2 else 1
        }

        return result

    def _parse_word(self, file_path: str, file_content: bytes) -> Dict:
        """解析Word文档"""
        result = {
            "format": DocumentFormat.WORD.value,
            "pages": [],
            "metadata": {}
        }

        temp_path = self._save_temp_file(file_path, file_content)

        try:
            doc = self._python_docx.Document(temp_path)

            # 提取文本
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            text = "\n".join(full_text)

            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # Word没有分页概念，全部作为一页
            result["pages"].append({
                "number": 1,
                "image": None,
                "text": text,
                "tables": tables
            })

            result["metadata"] = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }

        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

        return result

    def _parse_excel(self, file_path: str, file_content: bytes) -> Dict:
        """解析Excel文档"""
        result = {
            "format": DocumentFormat.EXCEL.value,
            "pages": [],
            "metadata": {}
        }

        temp_path = self._save_temp_file(file_path, file_content)

        try:
            wb = self._openpyxl.load_workbook(temp_path, data_only=True)

            for sheet_num, sheet in enumerate(wb.worksheets):
                # 提取数据
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    sheet_data.append(list(row))

                # 转换为文本
                text_rows = []
                for row in sheet_data:
                    text_rows.append("\t".join([str(cell) if cell is not None else "" for cell in row]))
                text = "\n".join(text_rows)

                result["pages"].append({
                    "number": sheet_num + 1,
                    "image": None,
                    "text": text,
                    "tables": [sheet_data] if sheet_data else []
                })

            result["metadata"] = {
                "sheet_count": len(wb.worksheets)
            }

        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

        return result

    def _extract_tables_from_page(self, page) -> List:
        """从PyMuPDF页面提取表格"""
        tables = []

        if self._has_pdfplumber:
            try:
                # 使用pdfplumber的表格检测
                import io
                import pdfplumber

                # 将页面转换为pdfplumber可用格式
                # 这里简化处理，实际可能需要更复杂的转换
                pass
            except Exception:
                pass

        return tables

    def _save_temp_file(self, file_path: str, file_content: bytes) -> str:
        """保存临时文件"""
        if file_content:
            ext = Path(file_path).suffix
            temp_path = os.path.join(self.temp_dir, f"temp_{os.urandom(8).hex()}{ext}")
            with open(temp_path, 'wb') as f:
                f.write(file_content)
            return temp_path
        else:
            return file_path

    def extract_text_from_page(self, page_data: Dict) -> str:
        """从页面数据提取文本"""
        return page_data.get("text", "")

    def extract_tables_from_page(self, page_data: Dict) -> List[Dict]:
        """从页面数据提取表格"""
        tables_data = page_data.get("tables", [])

        result = []
        for i, table in enumerate(tables_data):
            if isinstance(table, list):
                # 已经是二维数组格式
                result.append({
                    "index": i,
                    "headers": table[0] if table else [],
                    "rows": table[1:] if len(table) > 1 else [],
                    "row_count": len(table),
                    "col_count": len(table[0]) if table else 0
                })
            elif isinstance(table, dict):
                result.append(table)

        return result
