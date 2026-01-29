"""
非结构化文档 OCR 服务
Phase 1 P2: 数据感知汇聚 - OCR 处理

功能：
- PDF 文本提取
- 图片 OCR 识别
- 表格结构识别
- 版面分析
- 多语言支持（中英文）
"""

import base64
import io
import json
import logging
import os
import re
import tempfile
from datetime import datetime
from enum import Enum
from typing import Any, Dict, List, Optional, Tuple, Union
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# OCR 引擎配置
OCR_ENGINE = os.getenv("OCR_ENGINE", "auto")  # auto, tesseract, paddleocr, easyocr
TESSERACT_CMD = os.getenv("TESSERACT_CMD", "tesseract")
OCR_LANGUAGES = os.getenv("OCR_LANGUAGES", "chi_sim+eng")  # 默认中英文


class DocumentType(Enum):
    """文档类型"""
    PDF = "pdf"
    IMAGE = "image"
    WORD = "word"
    EXCEL = "excel"
    TEXT = "text"
    UNKNOWN = "unknown"


class OCREngine(Enum):
    """OCR 引擎类型"""
    TESSERACT = "tesseract"
    PADDLEOCR = "paddleocr"
    EASYOCR = "easyocr"
    AUTO = "auto"


@dataclass
class OCRResult:
    """OCR 识别结果"""
    text: str
    confidence: float = 0.0
    bounding_box: Optional[Dict[str, int]] = None  # {x, y, width, height}
    page_number: int = 1
    block_type: str = "text"  # text, table, figure, title


@dataclass
class TableCell:
    """表格单元格"""
    text: str
    row: int
    col: int
    row_span: int = 1
    col_span: int = 1


@dataclass
class ExtractedTable:
    """提取的表格"""
    cells: List[TableCell]
    num_rows: int
    num_cols: int
    page_number: int = 1

    def to_matrix(self) -> List[List[str]]:
        """转换为二维矩阵"""
        matrix = [["" for _ in range(self.num_cols)] for _ in range(self.num_rows)]
        for cell in self.cells:
            if 0 <= cell.row < self.num_rows and 0 <= cell.col < self.num_cols:
                matrix[cell.row][cell.col] = cell.text
        return matrix

    def to_markdown(self) -> str:
        """转换为 Markdown 表格"""
        matrix = self.to_matrix()
        if not matrix:
            return ""

        lines = []
        # 表头
        lines.append("| " + " | ".join(matrix[0]) + " |")
        lines.append("| " + " | ".join(["---"] * len(matrix[0])) + " |")
        # 数据行
        for row in matrix[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)


@dataclass
class DocumentExtractionResult:
    """文档提取结果"""
    text: str
    pages: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[ExtractedTable] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    ocr_results: List[OCRResult] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            "text": self.text,
            "page_count": len(self.pages),
            "pages": self.pages,
            "tables": [
                {
                    "page": t.page_number,
                    "rows": t.num_rows,
                    "cols": t.num_cols,
                    "data": t.to_matrix(),
                    "markdown": t.to_markdown(),
                }
                for t in self.tables
            ],
            "image_count": len(self.images),
            "images": self.images,
            "metadata": self.metadata,
            "char_count": len(self.text),
            "errors": self.errors,
        }


class OCRService:
    """OCR 服务"""

    def __init__(
        self,
        engine: str = None,
        languages: str = None,
    ):
        """
        初始化 OCR 服务

        Args:
            engine: OCR 引擎 (tesseract, paddleocr, easyocr, auto)
            languages: 语言设置
        """
        self.engine = engine or OCR_ENGINE
        self.languages = languages or OCR_LANGUAGES

        # 延迟初始化 OCR 引擎
        self._tesseract = None
        self._paddleocr = None
        self._easyocr = None

    def _get_tesseract(self):
        """获取 Tesseract OCR 实例"""
        if self._tesseract is None:
            try:
                import pytesseract
                pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
                self._tesseract = pytesseract
            except ImportError:
                raise ImportError("pytesseract is required. Install with: pip install pytesseract")
        return self._tesseract

    def _get_paddleocr(self):
        """获取 PaddleOCR 实例"""
        if self._paddleocr is None:
            try:
                from paddleocr import PaddleOCR
                self._paddleocr = PaddleOCR(
                    use_angle_cls=True,
                    lang='ch',
                    show_log=False,
                )
            except ImportError:
                raise ImportError("paddleocr is required. Install with: pip install paddleocr")
        return self._paddleocr

    def _get_easyocr(self):
        """获取 EasyOCR 实例"""
        if self._easyocr is None:
            try:
                import easyocr
                self._easyocr = easyocr.Reader(['ch_sim', 'en'])
            except ImportError:
                raise ImportError("easyocr is required. Install with: pip install easyocr")
        return self._easyocr

    def detect_document_type(self, file_path: str = None, content_type: str = None) -> DocumentType:
        """
        检测文档类型

        Args:
            file_path: 文件路径
            content_type: MIME 类型

        Returns:
            文档类型
        """
        if content_type:
            if 'pdf' in content_type.lower():
                return DocumentType.PDF
            elif 'image' in content_type.lower():
                return DocumentType.IMAGE
            elif 'word' in content_type.lower() or 'docx' in content_type.lower():
                return DocumentType.WORD
            elif 'excel' in content_type.lower() or 'xlsx' in content_type.lower():
                return DocumentType.EXCEL

        if file_path:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.pdf':
                return DocumentType.PDF
            elif ext in ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'):
                return DocumentType.IMAGE
            elif ext in ('.doc', '.docx'):
                return DocumentType.WORD
            elif ext in ('.xls', '.xlsx'):
                return DocumentType.EXCEL
            elif ext in ('.txt', '.md', '.csv'):
                return DocumentType.TEXT

        return DocumentType.UNKNOWN

    def extract_from_file(
        self,
        file_path: str,
        extract_tables: bool = True,
        extract_images: bool = False,
        ocr_images: bool = True,
    ) -> DocumentExtractionResult:
        """
        从文件提取内容

        Args:
            file_path: 文件路径
            extract_tables: 是否提取表格
            extract_images: 是否提取图片
            ocr_images: 是否对图片进行 OCR

        Returns:
            提取结果
        """
        doc_type = self.detect_document_type(file_path=file_path)

        if doc_type == DocumentType.PDF:
            return self._extract_from_pdf(file_path, extract_tables, extract_images, ocr_images)
        elif doc_type == DocumentType.IMAGE:
            return self._extract_from_image(file_path)
        elif doc_type == DocumentType.WORD:
            return self._extract_from_word(file_path)
        elif doc_type == DocumentType.TEXT:
            return self._extract_from_text(file_path)
        else:
            result = DocumentExtractionResult(text="")
            result.errors.append(f"Unsupported document type: {doc_type.value}")
            return result

    def extract_from_bytes(
        self,
        data: bytes,
        filename: str = None,
        content_type: str = None,
        extract_tables: bool = True,
        extract_images: bool = False,
        ocr_images: bool = True,
    ) -> DocumentExtractionResult:
        """
        从字节数据提取内容

        Args:
            data: 文件字节数据
            filename: 文件名
            content_type: MIME 类型
            extract_tables: 是否提取表格
            extract_images: 是否提取图片
            ocr_images: 是否对图片进行 OCR

        Returns:
            提取结果
        """
        # 保存到临时文件
        ext = ""
        if filename:
            ext = os.path.splitext(filename)[1]
        elif content_type:
            ext = self._mime_to_ext(content_type)

        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name

        try:
            return self.extract_from_file(tmp_path, extract_tables, extract_images, ocr_images)
        finally:
            os.unlink(tmp_path)

    def ocr_image(
        self,
        image_source: Union[str, bytes, "PIL.Image.Image"],
        engine: str = None,
    ) -> List[OCRResult]:
        """
        对图片进行 OCR

        Args:
            image_source: 图片路径、字节数据或 PIL Image
            engine: 指定 OCR 引擎

        Returns:
            OCR 结果列表
        """
        engine = engine or self.engine

        try:
            from PIL import Image
        except ImportError:
            raise ImportError("Pillow is required. Install with: pip install Pillow")

        # 转换为 PIL Image
        if isinstance(image_source, str):
            image = Image.open(image_source)
        elif isinstance(image_source, bytes):
            image = Image.open(io.BytesIO(image_source))
        else:
            image = image_source

        # 选择 OCR 引擎
        if engine == "auto":
            # 自动选择可用的引擎
            for eng in ["paddleocr", "tesseract", "easyocr"]:
                try:
                    return self._ocr_with_engine(image, eng)
                except ImportError:
                    continue
            raise ImportError("No OCR engine available. Install pytesseract, paddleocr, or easyocr.")
        else:
            return self._ocr_with_engine(image, engine)

    def _ocr_with_engine(self, image: "PIL.Image.Image", engine: str) -> List[OCRResult]:
        """使用指定引擎进行 OCR"""
        results = []

        if engine == "tesseract":
            pytesseract = self._get_tesseract()

            # 获取详细结果
            data = pytesseract.image_to_data(image, lang=self.languages, output_type=pytesseract.Output.DICT)

            for i in range(len(data['text'])):
                text = data['text'][i].strip()
                if text:
                    conf = float(data['conf'][i]) / 100.0 if data['conf'][i] != -1 else 0.0
                    results.append(OCRResult(
                        text=text,
                        confidence=conf,
                        bounding_box={
                            "x": data['left'][i],
                            "y": data['top'][i],
                            "width": data['width'][i],
                            "height": data['height'][i],
                        }
                    ))

        elif engine == "paddleocr":
            ocr = self._get_paddleocr()

            # 转换为 numpy 数组
            import numpy as np
            img_array = np.array(image)

            result = ocr.ocr(img_array, cls=True)

            if result and result[0]:
                for line in result[0]:
                    box, (text, confidence) = line
                    x_coords = [p[0] for p in box]
                    y_coords = [p[1] for p in box]
                    results.append(OCRResult(
                        text=text,
                        confidence=float(confidence),
                        bounding_box={
                            "x": int(min(x_coords)),
                            "y": int(min(y_coords)),
                            "width": int(max(x_coords) - min(x_coords)),
                            "height": int(max(y_coords) - min(y_coords)),
                        }
                    ))

        elif engine == "easyocr":
            reader = self._get_easyocr()

            import numpy as np
            img_array = np.array(image)

            result = reader.readtext(img_array)

            for box, text, confidence in result:
                x_coords = [p[0] for p in box]
                y_coords = [p[1] for p in box]
                results.append(OCRResult(
                    text=text,
                    confidence=float(confidence),
                    bounding_box={
                        "x": int(min(x_coords)),
                        "y": int(min(y_coords)),
                        "width": int(max(x_coords) - min(x_coords)),
                        "height": int(max(y_coords) - min(y_coords)),
                    }
                ))

        return results

    def _extract_from_pdf(
        self,
        file_path: str,
        extract_tables: bool = True,
        extract_images: bool = False,
        ocr_images: bool = True,
    ) -> DocumentExtractionResult:
        """从 PDF 提取内容"""
        result = DocumentExtractionResult(text="", metadata={"source": file_path})

        try:
            import fitz  # PyMuPDF
        except ImportError:
            result.errors.append("PyMuPDF is required for PDF processing. Install with: pip install PyMuPDF")
            return result

        try:
            doc = fitz.open(file_path)
            result.metadata.update({
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "created": doc.metadata.get("creationDate", ""),
            })

            all_text = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()

                page_info = {
                    "page_number": page_num + 1,
                    "text": page_text,
                    "width": page.rect.width,
                    "height": page.rect.height,
                }

                # 检查是否需要 OCR（文本很少的页面）
                if ocr_images and len(page_text.strip()) < 50:
                    # 将页面转换为图片进行 OCR
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")

                    try:
                        ocr_results = self.ocr_image(img_data)
                        ocr_text = " ".join([r.text for r in ocr_results])
                        if ocr_text:
                            page_text = ocr_text
                            page_info["text"] = ocr_text
                            page_info["ocr_applied"] = True
                            result.ocr_results.extend(ocr_results)
                    except Exception as e:
                        logger.warning(f"OCR failed for page {page_num + 1}: {e}")
                        result.errors.append(f"OCR failed for page {page_num + 1}: {str(e)}")

                # 提取表格
                if extract_tables:
                    tables = page.find_tables()
                    for table in tables:
                        cells = []
                        for row_idx, row in enumerate(table.extract()):
                            for col_idx, cell in enumerate(row):
                                if cell:
                                    cells.append(TableCell(
                                        text=str(cell),
                                        row=row_idx,
                                        col=col_idx,
                                    ))
                        if cells:
                            extracted_table = ExtractedTable(
                                cells=cells,
                                num_rows=len(table.extract()),
                                num_cols=len(table.extract()[0]) if table.extract() else 0,
                                page_number=page_num + 1,
                            )
                            result.tables.append(extracted_table)

                # 提取图片
                if extract_images:
                    images = page.get_images(full=True)
                    for img_idx, img in enumerate(images):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        result.images.append({
                            "page": page_num + 1,
                            "index": img_idx,
                            "width": base_image.get("width"),
                            "height": base_image.get("height"),
                            "ext": base_image.get("ext"),
                        })

                all_text.append(page_text)
                result.pages.append(page_info)

            result.text = "\n\n".join(all_text)
            doc.close()

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            result.errors.append(f"PDF extraction failed: {str(e)}")

        return result

    def _extract_from_image(self, file_path: str) -> DocumentExtractionResult:
        """从图片提取内容"""
        result = DocumentExtractionResult(text="", metadata={"source": file_path})

        try:
            ocr_results = self.ocr_image(file_path)
            result.ocr_results = ocr_results
            result.text = " ".join([r.text for r in ocr_results])

            # 添加图片信息
            try:
                from PIL import Image
                with Image.open(file_path) as img:
                    result.metadata.update({
                        "width": img.width,
                        "height": img.height,
                        "format": img.format,
                        "mode": img.mode,
                    })
            except Exception:
                pass

            result.pages.append({
                "page_number": 1,
                "text": result.text,
            })

        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            result.errors.append(f"Image OCR failed: {str(e)}")

        return result

    def _extract_from_word(self, file_path: str) -> DocumentExtractionResult:
        """从 Word 文档提取内容"""
        result = DocumentExtractionResult(text="", metadata={"source": file_path})

        try:
            import docx
        except ImportError:
            result.errors.append("python-docx is required for Word processing. Install with: pip install python-docx")
            return result

        try:
            doc = docx.Document(file_path)

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            result.text = "\n\n".join(paragraphs)

            # 提取表格
            for table_idx, table in enumerate(doc.tables):
                cells = []
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cells.append(TableCell(
                            text=cell.text,
                            row=row_idx,
                            col=col_idx,
                        ))

                if cells:
                    result.tables.append(ExtractedTable(
                        cells=cells,
                        num_rows=len(table.rows),
                        num_cols=len(table.columns),
                        page_number=1,
                    ))

            result.metadata.update({
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            })

            result.pages.append({
                "page_number": 1,
                "text": result.text,
            })

        except Exception as e:
            logger.error(f"Word extraction failed: {e}")
            result.errors.append(f"Word extraction failed: {str(e)}")

        return result

    def _extract_from_text(self, file_path: str) -> DocumentExtractionResult:
        """从文本文件提取内容"""
        result = DocumentExtractionResult(text="", metadata={"source": file_path})

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                result.text = f.read()

            result.metadata.update({
                "char_count": len(result.text),
                "line_count": result.text.count('\n') + 1,
            })

            result.pages.append({
                "page_number": 1,
                "text": result.text,
            })

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            result.errors.append(f"Text extraction failed: {str(e)}")

        return result

    def _mime_to_ext(self, content_type: str) -> str:
        """MIME 类型转扩展名"""
        mime_map = {
            "application/pdf": ".pdf",
            "image/jpeg": ".jpg",
            "image/png": ".png",
            "image/gif": ".gif",
            "image/webp": ".webp",
            "image/tiff": ".tiff",
            "application/msword": ".doc",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
            "application/vnd.ms-excel": ".xls",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
            "text/plain": ".txt",
        }
        return mime_map.get(content_type, "")

    def extract_structured_data(
        self,
        text: str,
        data_type: str = "auto",
    ) -> Dict[str, Any]:
        """
        从文本中提取结构化数据

        Args:
            text: 输入文本
            data_type: 数据类型 (auto, invoice, contract, id_card, etc.)

        Returns:
            结构化数据
        """
        result = {
            "type": data_type,
            "fields": {},
            "raw_text": text,
        }

        # 常见字段模式
        patterns = {
            # 发票信息
            "invoice_number": r"发票(?:号码|号)[：:]\s*(\w+)",
            "invoice_date": r"开票日期[：:]\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)",
            "invoice_amount": r"(?:金额|合计)[：:]\s*[¥￥]?\s*([\d,.]+)",

            # 身份证信息
            "id_number": r"(?:身份证号|公民身份号码)[：:]\s*(\d{17}[\dXx])",
            "name": r"姓名[：:]\s*([\u4e00-\u9fa5]{2,4})",
            "gender": r"性别[：:]\s*([男女])",
            "birth_date": r"(?:出生|出生日期)[：:]\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)",
            "address": r"(?:住址|地址)[：:]\s*([\u4e00-\u9fa5\d\-]+)",

            # 合同信息
            "contract_number": r"(?:合同编号|合同号)[：:]\s*([\w\-]+)",
            "contract_date": r"(?:签订日期|合同日期)[：:]\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)",
            "party_a": r"甲方[：:]\s*([\u4e00-\u9fa5\w]+)",
            "party_b": r"乙方[：:]\s*([\u4e00-\u9fa5\w]+)",

            # 通用字段
            "phone": r"(?:电话|手机|联系方式)[：:]\s*(1[3-9]\d{9})",
            "email": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
            "date": r"\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?",
            "amount": r"[¥￥$]\s*([\d,.]+)",
        }

        for field_name, pattern in patterns.items():
            matches = re.findall(pattern, text)
            if matches:
                result["fields"][field_name] = matches[0] if len(matches) == 1 else matches

        return result


# 创建全局实例
_ocr_service: Optional[OCRService] = None


def get_ocr_service() -> OCRService:
    """获取 OCR 服务单例"""
    global _ocr_service
    if _ocr_service is None:
        _ocr_service = OCRService()
    return _ocr_service
